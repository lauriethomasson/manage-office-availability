"""Per-file-type raw content extraction.

Each reader returns a dict with at least:
  text  - plain readable text (for LLM fallback / line-based rules)
  html  - raw HTML string, when the source is HTML/eml (for link-aware rules)
  links - list of (visible_text, href) tuples in document order, when available
  tables - list of tables, each a list of rows, each row a list of cell strings
  row_links - .xlsx/.xls only: list of {"row_text", "links"} per source row
    that has a real hyperlink (see _extract_xlsx_row_links) — pandas'
    own cell-value read (used for text/tables above) discards hyperlinks
    entirely, so this is the only place they're recoverable at all

Readers raise ValueError with a human-readable message on failure — the
caller (pipeline) turns that into a per-file error in the results summary
instead of crashing the whole batch.
"""
import csv
import email
import re
from datetime import date as _date
from email import policy
from io import StringIO
from pathlib import Path

from bs4 import BeautifulSoup

PDF_DATE_RE = re.compile(r"D:(\d{4})(\d{2})(\d{2})")


def read_file(path):
    ext = Path(path).suffix.lower()
    try:
        if ext == ".pdf":
            return _read_pdf(path)
        if ext == ".docx":
            return _read_docx(path)
        if ext in (".xlsx", ".xls"):
            return _read_xlsx(path)
        if ext == ".csv":
            return _read_csv(path)
        if ext == ".eml":
            return _read_eml(path)
        if ext in (".html", ".htm"):
            return _read_html(path)
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Could not read {ext} file: {e}")
    raise ValueError(f"Unsupported file type: {ext}")


def _clean_pdf_cell(cell):
    """PDF table cells often wrap long text across multiple physical lines;
    join them back into prose, without adding a stray space after a
    hyphenated line break (e.g. "All-\nInclusive" -> "All-Inclusive")."""
    if not cell:
        return ""
    out = ""
    for line in cell.replace("\r", "").split("\n"):
        line = line.strip()
        if not line:
            continue
        # Only treat a trailing "-" as a hyphenated word-break (join with no
        # space) when it directly follows a letter, e.g. "All-"; a standalone
        # dash like "CAT A -" is punctuation and should keep its space.
        hyphenated = len(out) >= 2 and out.endswith("-") and out[-2].isalnum()
        out = out + line if hyphenated else (out + " " + line).strip()
    return out


def _empty(**overrides):
    base = {
        "text": "",
        "html": "",
        "links": [],
        "html_items": [],
        "tables": [],
        "file_date": None,
        "pages_text": [],
        "size_warning": None,
        "row_links": [],
    }
    base.update(overrides)
    return base


def _parse_pdf_date(raw):
    """PDF metadata dates look like "D:20260630101530+01'00'" — pull out
    just the YYYY-MM-DD, ignoring the time/timezone. Returns None if `raw`
    isn't present or doesn't match the expected format."""
    if not raw:
        return None
    m = PDF_DATE_RE.match(raw)
    if not m:
        return None
    year, month, day = (int(g) for g in m.groups())
    try:
        return _date(year, month, day).isoformat()
    except ValueError:
        return None


# A backstop against genuinely pathological documents, not a tuning knob
# for normal brochures — the largest real source seen so far (Crown
# Estate, 20 pages) is nowhere near this. Checked before any per-page
# text/table/image extraction is attempted, so a document that would
# take an unreasonable amount of time/memory to process on the free
# tier's constraints fails fast with a clear reason instead of a long
# processing attempt that then still risks a timeout or OOM.
MAX_PDF_PAGES = 300

# Distinct from MAX_PDF_PAGES above: this is what's actually been proven
# end-to-end (real HTTP request -> LLM extraction -> geocoding -> image/
# floor-plan extraction -> spreadsheet), repeatedly, with verified-correct
# output — not a guess, and not the same claim as "hasn't hit the hard
# 300-page ceiling yet". The largest real file exercised through the full
# pipeline so far: Crown Estate, 20 pages / ~4.3MB. TESTED_MAX_PDF_BYTES
# sits a little above that exact figure so the known-good reference file
# itself doesn't spuriously trigger this. A file beyond these thresholds
# (but still under MAX_PDF_PAGES) is a soft *warning*, not an error — it
# may well process fine, page-by-page memory bounding (extraction/
# pdf_images.py) doesn't stop applying just because a file is bigger than
# what's been tested — this only tells whoever's running the batch that
# this exact size combination hasn't been specifically verified yet.
TESTED_MAX_PDF_PAGES = 20
TESTED_MAX_PDF_BYTES = 4.5 * 1024 * 1024


def _read_pdf(path):
    import pdfplumber

    text_parts = []
    tables = []
    file_date = None
    try:
        with pdfplumber.open(path) as pdf:
            if not pdf.pages:
                raise ValueError("PDF has no pages")
            if len(pdf.pages) > MAX_PDF_PAGES:
                raise ValueError(
                    f"PDF has {len(pdf.pages)} pages, exceeding the {MAX_PDF_PAGES}-page limit for processing on "
                    "this plan — split it into smaller files and process them separately"
                )
            metadata = pdf.metadata or {}
            # Prefer CreationDate (closer to "when this was actually put
            # together/sent") over ModDate, but take whichever parses.
            file_date = _parse_pdf_date(metadata.get("CreationDate")) or _parse_pdf_date(metadata.get("ModDate"))
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                # Always append, even "" for a text-less page — pages_text
                # below must stay index-aligned with the PDF's own 0-indexed
                # page numbers (extraction.pdf_images iterates every page
                # unconditionally), not just the pages that had extractable
                # text.
                text_parts.append(page_text)
                for t in page.extract_tables() or []:
                    cleaned = [[_clean_pdf_cell(c) for c in row] for row in t]
                    if cleaned:
                        tables.append(cleaned)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {e}")

    text = "\n".join(text_parts).strip()
    if not text and not tables:
        raise ValueError("No extractable text found in PDF (it may be a scanned image)")

    page_count = len(text_parts)
    try:
        size_bytes = Path(path).stat().st_size
    except OSError:
        size_bytes = 0
    size_warning = None
    if page_count > TESTED_MAX_PDF_PAGES or size_bytes > TESTED_MAX_PDF_BYTES:
        size_warning = (
            f"This PDF ({page_count} pages, {size_bytes / 1024 / 1024:.1f}MB) is larger than what's "
            f"been fully tested end-to-end (~{TESTED_MAX_PDF_PAGES} pages / ~4MB) — it may well process "
            "fine, it just hasn't been specifically verified at this size yet."
        )

    # Per-page text (not just the joined whole), so a downstream Floor
    # Plan/High Res Images enrichment step can tell which of the source
    # PDF's own pages a given extracted listing actually came from — see
    # extraction.pdf_images.find_matching_pages.
    return _empty(text=text, tables=tables, file_date=file_date, pages_text=text_parts, size_warning=size_warning)


def _read_docx(path):
    import docx

    try:
        d = docx.Document(path)
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {e}")

    text_parts = [p.text for p in d.paragraphs if p.text.strip()]
    tables = []
    for t in d.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in t.rows]
        if rows:
            tables.append(rows)
            for row in rows:
                text_parts.append(" | ".join(row))

    text = "\n".join(text_parts).strip()
    if not text:
        raise ValueError("No text found in DOCX")

    file_date = None
    for dt in (d.core_properties.created, d.core_properties.modified):
        if dt:
            file_date = dt.date().isoformat()
            break
    return _empty(text=text, tables=tables, file_date=file_date)


def _read_xlsx(path):
    import pandas as pd

    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    except Exception as e:
        raise ValueError(f"Failed to parse spreadsheet: {e}")

    tables = []
    text_parts = []
    for name, df in sheets.items():
        df = df.fillna("")
        header = [str(c) for c in df.columns]
        rows = [header] + df.astype(str).values.tolist()
        tables.append(rows)
        text_parts.append(f"[Sheet: {name}]")
        for row in rows:
            text_parts.append(" | ".join(row))

    if not tables or all(len(t) <= 1 for t in tables):
        raise ValueError("Spreadsheet appears to have no data rows")

    # Best-effort: real per-row hyperlink data pandas discards entirely
    # (see _extract_xlsx_row_links's own docstring) but isn't essential to
    # basic extraction — a workbook openpyxl can't open for any reason
    # (corruption, an unusual format variant) shouldn't fail the whole
    # file when pandas already read it successfully above.
    try:
        row_links = _extract_xlsx_row_links(path)
    except Exception:
        row_links = []

    return _empty(text="\n".join(text_parts), tables=tables, row_links=row_links)


def _extract_xlsx_row_links(path):
    """Returns a list of {"row_text": str, "links": [(display_text, url), ...]}
    — one entry per source row with at least one real hyperlink, across
    every sheet, in workbook order.

    pandas.read_excel above (used for text/tables) reads cell VALUES only
    and silently discards hyperlinks entirely. Confirmed via a real UNION
    source (2026-07): its own "Brochure" column links every row to a real
    box.com brochure/floor-plan URL through a hyperlink on a generic
    "CLICK HERE"/"Landlord Brochure"/"FLOOR PLAN" display cell, never the
    URL itself as visible text — so nothing in the LLM's own plain-text
    prompt input (built from exactly the same pandas-read values) could
    ever recover it, and Brochure PDF/Floor Plan came back blank for
    every row despite real links existing in the source.

    row_text mirrors the same row's own dumped text (the " | "-joined
    cell values already used for the LLM's text input above) so a later
    enrichment step (extraction.xlsx_links) can match it back to a
    specific extracted listing by Building-name substring search — the
    same matching principle already used for a PDF (extraction.
    pdf_images, matched by page) or an .eml (extraction.html_images,
    matched by alt text), just keyed by row here instead."""
    from openpyxl import load_workbook

    wb = load_workbook(path, data_only=True)
    row_links = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            links = [
                (str(cell.value).strip() if cell.value is not None else "", cell.hyperlink.target)
                for cell in row
                if cell.hyperlink is not None and cell.hyperlink.target
            ]
            if not links:
                continue
            row_text = " | ".join(str(cell.value) for cell in row if cell.value not in (None, ""))
            row_links.append({"row_text": row_text, "links": links})
    return row_links


def _read_csv(path):
    try:
        with open(path, newline="", encoding="utf-8-sig") as f:
            rows = list(csv.reader(f))
    except UnicodeDecodeError:
        with open(path, newline="", encoding="latin-1") as f:
            rows = list(csv.reader(f))
    except Exception as e:
        raise ValueError(f"Failed to parse CSV: {e}")

    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        raise ValueError("CSV file is empty")

    buf = StringIO()
    for row in rows:
        buf.write(" | ".join(row) + "\n")
    return _empty(text=buf.getvalue(), tables=[rows])


def _read_eml(path):
    try:
        with open(path, "rb") as f:
            msg = email.message_from_binary_file(f, policy=policy.default)
    except Exception as e:
        raise ValueError(f"Failed to parse EML: {e}")

    html_body = None
    text_body = None
    for part in msg.walk():
        ctype = part.get_content_type()
        try:
            content = part.get_content()
        except Exception:
            continue
        if ctype == "text/html" and html_body is None:
            html_body = content
        elif ctype == "text/plain" and text_body is None:
            text_body = content

    if html_body:
        result = _parse_html_string(html_body)
    elif text_body:
        result = _empty(text=text_body.strip())
    else:
        raise ValueError("EML has no readable text or HTML body")

    result["subject"] = msg.get("Subject", "")
    result["sender"] = msg.get("From", "")
    result["date"] = msg.get("Date", "")
    if not result["text"].strip():
        raise ValueError("EML body is empty after parsing")
    return result


def _read_html(path):
    try:
        html = Path(path).read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        raise ValueError(f"Failed to read HTML file: {e}")
    result = _parse_html_string(html)
    if not result["text"].strip():
        raise ValueError("HTML file has no visible text")
    return result


def _parse_html_string(html):
    soup = BeautifulSoup(html, "lxml")
    # Some senders hard-wrap text with literal newlines *inside* a single
    # text node (e.g. "View\n property", "Shared\n garden..."), which is a
    # source-formatting artifact, not a real line break. Collapse internal
    # whitespace per node so each node becomes exactly one clean line, then
    # join nodes with "\n" — this keeps genuine element-to-element breaks.
    lines = [re.sub(r"\s+", " ", s).strip() for s in soup.stripped_strings]
    text = "\n".join(l for l in lines if l)
    links = [
        (re.sub(r"\s+", " ", a.get_text(" ", strip=True)).strip(), a.get("href", ""))
        for a in soup.find_all("a")
        if a.get("href")
    ]
    # Interleaved (kind, text_or_alt, href_or_src) walk of <a>/<img> in true
    # document order — lets a link-grouping rule (extraction.rules.knotel)
    # correlate a listing's own photo (an <img>, not a link) to whichever
    # group of buttons immediately follows it in the same card, something
    # `links` above can't do since it only ever tracks anchors and drops
    # images entirely. Confirmed empirically (Knotel) that a listing's own
    # photo is a genuinely separate <img alt="... featured image"> — not a
    # decorative logo/icon — appearing right before that listing's button
    # row in the DOM, so a source's own layout convention decides "kind" via
    # alt text here, not this function.
    html_items = []
    for el in soup.find_all(["a", "img"]):
        if el.name == "a":
            href = el.get("href", "")
            if not href:
                continue
            link_text = re.sub(r"\s+", " ", el.get_text(" ", strip=True)).strip()
            html_items.append(("link", link_text, href))
        else:
            src = el.get("src", "")
            alt = el.get("alt", "")
            # Confirmed empirically (MetSpace) that a source's own real,
            # per-listing photos can legitimately have alt="" — requiring
            # alt to be non-empty here (as originally written, tuned
            # against Knotel's "X Floor featured image" convention) silently
            # dropped every one of them before a rule ever saw them. Keep
            # any image with a src at all; a specific rule (extraction.
            # rules.knotel filters by alt text, extraction.rules.metspace
            # by source domain) decides what's a real content photo vs.
            # decorative — this layer shouldn't guess that generically.
            if src:
                html_items.append(("image", alt, src))
    tables = []
    for t in soup.find_all("table"):
        rows = []
        for tr in t.find_all("tr"):
            cells = [c.get_text(" ", strip=True) for c in tr.find_all(["td", "th"])]
            if any(c for c in cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
    return _empty(text=text, html=html, links=links, html_items=html_items, tables=tables)
